from docx import Document
import os
filename: str = r"assets\actFile.docx"



def EditInFile(User, Director, RecievDate, ReturnDate, type, item, act,user_status,dir_status):
    doc = Document(filename)
    pr = doc.paragraphs

    for i in range(len(pr)):
        if "[FIRSTDATE]" in pr[i].text:
            pr[i].text = pr[i].text.replace("[FIRSTDATE]", RecievDate)
        if "[USER]" in pr[i].text:
            pr[i].text = pr[i].text.replace("[USER]", User)
        if "[DIR]" in pr[i].text:
            pr[i].text = pr[i].text.replace("[DIR]", Director)
        if "[LASTDATE]" in pr[i].text:
            pr[i].text = pr[i].text.replace("[LASTDATE]", ReturnDate)
        if "[ACT]" in pr[i].text:
            pr[i].text = pr[i].text.replace("[ACT]", act)
        if "[USERSTATUS]" in pr[i].text:
            pr[i].text = pr[i].text.replace("[USERSTATUS]", user_status)
        if "[DIRSTATUS]" in pr[i].text:
            pr[i].text = pr[i].text.replace("[DIRSTATUS]", dir_status)


    table = doc.tables[0]
    if type == "Computer":
        table.cell(1, 1).text = "Ноутбук"
        table.cell(2, 1).text = "Зарядник"
    else:
        table.cell(1, 1).text = "Планшет"
        row = table.rows[2]
        row._element.getparent().remove(row._element)
    table.cell(1, 2).text = item[2] # models
    table.cell(1, 3).text = item[3] # SN
    table.cell(1, 4).text = item[1] # Code



    newfilename =  rf"assets\{act}_ACT_{User}.docx"
    doc.save(newfilename)
    os.startfile(newfilename)
    print("Done")


def open_file():
    if os.path.exists(filename):
        os.startfile(filename)
    else:
        print("File does not exist at the given path")

def open_folder():
    os.startfile("assets")